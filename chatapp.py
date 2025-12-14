
import streamlit as st
import networkx as nx
import matplotlib.pyplot as plt
import re
from collections import Counter
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
import io
from fpdf import FPDF
import docx

load_dotenv()
if os.getenv("GOOGLE_API_KEY"):
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def ensure_api_key():
    key = st.session_state.get("GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if not key:
        return False
    os.environ["GOOGLE_API_KEY"] = key
    genai.configure(api_key=key)
    return True

# Lightweight local embedding to avoid external quotas
class SimpleHashEmbeddings:
    def __init__(self, dim: int = 512):
        self.dim = dim

    def _embed(self, text: str):
        vec = [0.0] * self.dim
        for token in (text or "").lower().split():
            idx = (hash(token) % self.dim)
            vec[idx] += 1.0
        norm = sum(v*v for v in vec) ** 0.5
        if norm > 0:
            vec = [v / norm for v in vec]
        return vec

    def embed_documents(self, texts):
        return [self._embed(t) for t in texts]

    def embed_query(self, text):
        return self._embed(text)

def get_embeddings_provider():
    provider = st.session_state.get("EMBEDDING_PROVIDER")
    if provider == "Local (No-Quota)":
        return SimpleHashEmbeddings()
    if ensure_api_key():
        return GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    return SimpleHashEmbeddings()



def extract_table_of_contents(text):
    # Simple TOC extraction: look for lines starting with numbers or section keywords
    import re
    toc = re.findall(r'^(?:\d+\.|Section|Chapter|Part)\s.*', text, re.MULTILINE)
    return toc if toc else ["No TOC found."]

def extract_summary(text):
    # Use LLM or simple heuristic for summary (pseudo, replace with LLM if available)
    summary = text[:500] + "..." if len(text) > 500 else text
    return summary

def extract_tags_entities(text):
    """
    Use a simple regex and LLM to extract topics, entities, and tags from text.
    For demo, use regex for capitalized words and LLM for advanced extraction.
    """
    # Basic entity extraction (capitalized words)
    words = re.findall(r'\b[A-Z][a-zA-Z]{2,}\b', text)
    common_words = set(['The', 'This', 'That', 'With', 'From', 'Page', 'PDF', 'Document'])
    entities = [w for w in words if w not in common_words]
    # Frequency-based tags
    tags = [w for w, c in Counter(entities).most_common(10)]
    # LLM-based extraction (pseudo, replace with actual call if needed)
    # You can use GoogleGenerativeAIEmbeddings or ChatGoogleGenerativeAI for advanced extraction
    return list(set(tags)), list(set(entities))

def get_pdf_text(pdf_docs):
    text = ""
    pdf_texts = []
    for pdf in pdf_docs:
        fname = pdf.name if hasattr(pdf, 'name') else str(pdf)
        if fname.lower().endswith('.pdf'):
            pdf_reader = PdfReader(pdf)
            doc_text = ""
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                doc_text += page_text
                pdf_texts.append({
                    "filename": fname,
                    "page": i+1,
                    "text": page_text,
                    "type": "pdf"
                })
            text += doc_text
        elif fname.lower().endswith('.txt'):
            content = pdf.read().decode('utf-8')
            pdf_texts.append({
                "filename": fname,
                "page": 1,
                "text": content,
                "type": "txt"
            })
            text += content
        elif fname.lower().endswith('.docx'):
            doc = docx.Document(pdf)
            doc_text = "\n".join([p.text for p in doc.paragraphs])
            pdf_texts.append({
                "filename": fname,
                "page": 1,
                "text": doc_text,
                "type": "docx"
            })
            text += doc_text
    # Extract summary, TOC, tags/entities from all text
    summary = extract_summary(text)
    toc = extract_table_of_contents(text)
    tags, entities = extract_tags_entities(text)
    return text, pdf_texts, tags, entities, summary, toc




def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=50000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embedding_model_name = "models/embedding-001"
    embeddings = get_embeddings_provider()
    try:
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
    except Exception as e:
        import streamlit as st
        st.error(f"Error embedding content: {e}")
        if not isinstance(embeddings, SimpleHashEmbeddings):
            st.info("Try switching Embedding Provider to 'Local (No-Quota)' in the sidebar.")
        return None


def get_conversational_chain():

    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    chat_model_name = "models/gemini-1.5-flash-latest"
    model = ChatGoogleGenerativeAI(model=chat_model_name, temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return (model, prompt)




def user_input(user_question, chat_history, pdf_texts):
    embedding_model_name = "models/embedding-001"
    embeddings = get_embeddings_provider()
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    model, prompt = get_conversational_chain()
    formatted = prompt.format(context="\n".join([d.page_content for d in docs]), question=user_question)
    answer = model.invoke(formatted)

    # Find source reference (filename and page) and preview
    source_info = []
    preview_text = ""
    for doc in docs:
        for pdf in pdf_texts:
            if pdf["text"] and pdf["text"].strip()[:30] in doc.page_content[:100]:
                source_info.append(f"{pdf['filename']} (Page {pdf['page']})")
                preview_text = pdf["text"][:300]  # Show first 300 chars as preview
                break
    source_info = list(set(source_info))

    chat_history.append({
        "question": user_question,
        "answer": answer,
        "source": ", ".join(source_info) if source_info else "Unknown",
        "preview": preview_text
    })
    return answer, chat_history





def main():
    st.set_page_config("Multi PDF Chatbot", page_icon=":scroll:")
    st.header("Multi-PDF's 📚 - Chat Agent 🤖 ")

    # Session state for chat history and PDF text
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "pdf_texts" not in st.session_state:
        st.session_state.pdf_texts = []
    if "raw_text" not in st.session_state:
        st.session_state.raw_text = ""
    if "tags" not in st.session_state:
        st.session_state.tags = []
    if "entities" not in st.session_state:
        st.session_state.entities = []
    if "summary" not in st.session_state:
        st.session_state.summary = ""
    if "toc" not in st.session_state:
        st.session_state.toc = []


    # Semantic Search & Contextual Recommendations
    st.subheader("🔍 Semantic Search Across All PDFs")
    search_query = st.text_input("Ask a question or search semantically:")
    filter_filename = st.selectbox("Filter by file", ["All"] + list(set([pdf["filename"] for pdf in st.session_state.pdf_texts])) if st.session_state.pdf_texts else ["All"])
    filter_page = st.number_input("Filter by page (0 = all)", min_value=0, value=0, step=1)
    if search_query and st.session_state.pdf_texts:
        # Use vector similarity search
        embedding_model_name = "models/embedding-001"
        embeddings = GoogleGenerativeAIEmbeddings(model=embedding_model_name)
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(search_query, k=5)
        results = []
        for doc in docs:
            for pdf in st.session_state.pdf_texts:
                if pdf["text"] and pdf["text"].strip()[:30] in doc.page_content[:100]:
                    if (filter_filename == "All" or pdf["filename"] == filter_filename) and (filter_page == 0 or pdf["page"] == filter_page):
                        snippet = pdf["text"][:300]
                        results.append(f"{pdf['filename']} (Page {pdf['page']}): ...{snippet}...")
        if results:
            st.info("\n".join(results))
        else:
            st.warning("No semantic results found.")
        # Contextual Recommendations: Suggest related questions
        st.markdown("**Related Questions & Follow-ups:**")
        # Use LLM to generate suggestions (pseudo, replace with actual call if needed)
        follow_ups = [f"What are the main topics in {filter_filename if filter_filename != 'All' else 'these documents'}?",
                     f"Summarize page {filter_page} of {filter_filename}" if filter_page != 0 else "Summarize this document.",
                     "Show entities and relationships."]
        for q in follow_ups:
            st.button(q, key=f"followup_{q}")



    # Multi-language support
    st.subheader("🌐 Language Settings")
    lang = st.selectbox("Select language for Q&A", ["English", "Hindi", "French", "German", "Spanish"])

    # --- Tabs for main app sections ---
    tabs = st.tabs(["Document Analysis", "AI Chatbot (One-to-One Conversation)", "Auto Insights & Anomaly Detection"])

    # Tab 2: AI Chatbot (One-to-One Conversation)
    with tabs[1]:
        st.markdown("""
        <style>
        .chat-bubble-user {background:linear-gradient(90deg,#232526,#414345);color:#fff;padding:10px;border-radius:10px;margin-bottom:5px;max-width:80%;margin-left:auto;text-align:right;box-shadow:0 2px 8px #0002;}
        .chat-bubble-ai {background:linear-gradient(90deg,#141e30,#243b55);color:#fff;padding:10px;border-radius:10px;margin-bottom:5px;max-width:80%;margin-right:auto;text-align:left;box-shadow:0 2px 8px #0002;}
        .chat-avatar {width:32px;height:32px;border-radius:50%;display:inline-block;vertical-align:middle;margin-right:8px;}
        .chat-row {display:flex;align-items:flex-end;}
        .chat-actions {margin-top:8px;}
        .sticky-input {position:fixed;bottom:0;left:0;width:100%;background:#232526;padding:10px;z-index:100;}
        </style>
        """, unsafe_allow_html=True)
        st.subheader("💬 AI Chatbot (One-to-One Conversation)")
        # Optional: Upload new document directly from chat tab
        st.markdown("**Upload a new document (PDF, DOCX, TXT):**")
        chat_docs = st.file_uploader("Upload here to add to knowledge base", accept_multiple_files=True, type=["pdf", "txt", "docx"], key="chat_doc_upload")
        if chat_docs:
            with st.spinner("Processing new document(s)..."):
                raw_text, pdf_texts, tags, entities, summary, toc = get_pdf_text(chat_docs)
                st.session_state.pdf_texts += pdf_texts
                st.session_state.raw_text += raw_text
                st.session_state.tags += tags
                st.session_state.entities += entities
                st.session_state.summary += "\n" + summary
                st.session_state.toc += toc
                st.success("Document(s) added!")
        # Chat input sticky at bottom
        st.markdown('<div class="sticky-input">', unsafe_allow_html=True)
        user_question = st.text_input("Type your question for the AI:", key="chat_input")
        st.markdown('</div>', unsafe_allow_html=True)
        # Quick action buttons for follow-up questions
        follow_ups = ["Summarize all documents", "List key topics/entities", "Show table of contents", "Show document relationships"]
        st.markdown('<div class="chat-actions">', unsafe_allow_html=True)
        for q in follow_ups:
            if st.button(q, key=f"followup_{q}"):
                st.session_state["chat_input"] = q
                st.experimental_rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        # Typing indicator
        if st.session_state.get("is_typing", False):
            st.markdown("<span style='color:#ffd700;font-weight:bold;'>AI is typing...</span>", unsafe_allow_html=True)
        # Chat logic
        if user_question and st.session_state.pdf_texts:
            st.session_state["is_typing"] = True
            system_prompt = f"""
You are a world-class, highly accurate, step-by-step AI assistant. The user has uploaded multiple documents (PDF, DOCX, TXT). You have access to:
- Document summaries: {st.session_state.summary}
- Table of contents: {st.session_state.toc}
- Key topics/entities: {', '.join(st.session_state.entities)}
- Raw text: {st.session_state.raw_text[:1000]}... (truncated)

Always:
- Use all available document knowledge to answer questions as accurately and thoroughly as possible.
- If code or references are needed, provide Python code blocks and explain them.
- If the question is ambiguous, ask clarifying questions.
- If you don't know, say so honestly.
- Format your answer with clear sections, bullet points, and tables if helpful.
- Respond in a friendly, conversational, and professional tone.
"""
            model, prompt = get_conversational_chain()
            class Doc:
                def __init__(self, page_content, metadata=None):
                    self.page_content = page_content
                    self.metadata = metadata or {}
            docs = [Doc(system_prompt + '\n' + st.session_state.raw_text)]
            formatted = prompt.format(context="\n".join([d.page_content for d in docs]), question=user_question)
            answer = model.invoke(formatted)
            st.session_state["is_typing"] = False
            if "chat_history" not in st.session_state:
                st.session_state.chat_history = []
            st.session_state.chat_history.append({
                "question": user_question,
                "answer": answer,
                "source": "All Uploaded Documents",
                "preview": st.session_state.raw_text[:300]
            })
        # Enhanced chat history display
        import hashlib
        if st.session_state.chat_history:
            st.markdown("---")
            st.subheader("🗂️ Chat History")
            for i, chat in enumerate(st.session_state.chat_history[::-1]):
                question_hash = hashlib.md5(chat['question'].encode()).hexdigest()[:8]
                unique_key = f"askagain_{i}_{question_hash}"
                # Chat bubble UI
                avatar = "<img src='https://avatars.githubusercontent.com/u/9919?s=40&v=4' class='chat-avatar'/>" if i % 2 == 0 else "<img src='https://avatars.githubusercontent.com/u/583231?s=40&v=4' class='chat-avatar'/>"
                bubble_class = "chat-bubble-user" if i % 2 == 0 else "chat-bubble-ai"
                ts = chat.get('timestamp', f"{i+1}")
                st.markdown(f"<div class='chat-row'><div class='{bubble_class}'>{avatar}<b>{'You' if i % 2 == 0 else 'AI'}</b> <span style='font-size:10px;color:#bbb;'>{ts}</span><br>{chat['question'] if i % 2 == 0 else chat['answer']}</div></div>", unsafe_allow_html=True)
                # Markdown rendering for AI answers
                if i % 2 != 0:
                    st.markdown(chat['answer'], unsafe_allow_html=True)
                # Feedback buttons
                if i % 2 != 0:
                    colf1, colf2 = st.columns([1,1])
                    with colf1:
                        st.button("👍", key=f"thumbsup_{unique_key}")
                    with colf2:
                        st.button("👎", key=f"thumbsdown_{unique_key}")
                # Ask Again button
                if st.button(f"Ask Again: {chat['question']}", key=unique_key):
                    st.session_state["chat_input"] = chat['question']
                    st.experimental_rerun()
                # PDF page preview with highlight (for PDFs only)
                if chat['source'] != "Unknown" and "pdf" in chat['source'].lower():
                    src = chat['source'].split(' (Page ')
                    if len(src) == 2:
                        fname = src[0]
                        try:
                            page_num = int(src[1].replace(")", ""))
                            for pdf in st.session_state.pdf_texts:
                                if pdf['filename'] == fname and pdf['page'] == page_num and pdf['type'] == 'pdf':
                                    st.markdown(f"**PDF Page Preview:** {fname} (Page {page_num})")
                                    preview = pdf['text'].replace(chat['answer'], f"**{chat['answer']}**") if chat['answer'] in pdf['text'] else pdf['text']
                                    st.code(preview[:1000], language='text')
                        except Exception:
                            pass
                st.write("")
            # Download chat history as TXT or PDF
            chat_str = "\n\n".join([f"Q: {c['question']}\nA: {c['answer']}\nSource: {c['source']}" for c in st.session_state.chat_history])
            st.download_button("Download Chat History (TXT)", chat_str, file_name="chat_history.txt")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for c in st.session_state.chat_history:
                q = c['question'].encode('latin1', 'replace').decode('latin1')
                a = c['answer'].encode('latin1', 'replace').decode('latin1')
                s = c['source'].encode('latin1', 'replace').decode('latin1')
                pdf.multi_cell(0, 10, f"Q: {q}\nA: {a}\nSource: {s}\n\n")
            pdf_bytes = pdf.output(dest='S').encode('latin1')
            st.download_button("Download Chat History (PDF)", pdf_bytes, file_name="chat_history.pdf")

    # Chat history display with clickable questions, source preview, and PDF preview
    if st.session_state.chat_history:
        st.markdown("---")
        st.subheader("🗂️ Chat History")
        for i, chat in enumerate(st.session_state.chat_history[::-1]):
            if st.button(f"Ask Again: {chat['question']}", key=f"askagain_{i}"):
                st.session_state["chat_input"] = chat['question']
                st.experimental_rerun()
            st.markdown(f"**Q{i+1}:** {chat['question']}")
            st.markdown(f"**A{i+1}:** {chat['answer']}")
            st.caption(f"Source: {chat['source']}")
            if chat['preview']:
                st.code(chat['preview'], language='text')
            # PDF page preview with highlight (for PDFs only)
            if chat['source'] != "Unknown" and "pdf" in chat['source'].lower():
                src = chat['source'].split(' (Page ')
                if len(src) == 2:
                    fname = src[0]
                    try:
                        page_num = int(src[1].replace(")", ""))
                        for pdf in st.session_state.pdf_texts:
                            if pdf['filename'] == fname and pdf['page'] == page_num and pdf['type'] == 'pdf':
                                st.markdown(f"**PDF Page Preview:** {fname} (Page {page_num})")
                                # Highlight answer in preview (simple bold)
                                preview = pdf['text'].replace(chat['answer'], f"**{chat['answer']}**") if chat['answer'] in pdf['text'] else pdf['text']
                                st.code(preview[:1000], language='text')
                    except Exception:
                        pass
            st.write("")


    # Download chat history as TXT or PDF
    if st.session_state.chat_history:
        chat_str = "\n\n".join([f"Q: {c['question']}\nA: {c['answer']}\nSource: {c['source']}" for c in st.session_state.chat_history])
        st.download_button("Download Chat History (TXT)", chat_str, file_name="chat_history.txt")
        # PDF export with safe encoding
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for c in st.session_state.chat_history:
            # Replace non-latin1 characters with a safe replacement
            q = c['question'].encode('latin1', 'replace').decode('latin1')
            a = c['answer'].encode('latin1', 'replace').decode('latin1')
            s = c['source'].encode('latin1', 'replace').decode('latin1')
            pdf.multi_cell(0, 10, f"Q: {q}\nA: {a}\nSource: {s}\n\n")
        pdf_bytes = pdf.output(dest='S').encode('latin1')
        st.download_button("Download Chat History (PDF)", pdf_bytes, file_name="chat_history.pdf")


    # Selective Summarization feature
    st.subheader("📝 Summarize a Document or Page")
    if st.session_state.pdf_texts:
        filenames = list(set([pdf["filename"] for pdf in st.session_state.pdf_texts]))
        selected_file = st.selectbox("Select file to summarize", ["All"] + filenames)
        selected_page = st.number_input("Page to summarize (0 = all)", min_value=0, value=0, step=1)
        if st.button("Summarize Selection"):
            chat_model_name = "models/gemini-1.5-flash-latest"
            model = ChatGoogleGenerativeAI(model=chat_model_name, temperature=0.3)
            prompt = PromptTemplate(template="Summarize the following text as clearly as possible:\n{text}", input_variables=["text"])
            # Use direct prompt formatting with the chat model
            if selected_file == "All":
                text_to_summarize = st.session_state.raw_text
            else:
                text_to_summarize = ""
                for pdf in st.session_state.pdf_texts:
                    if pdf["filename"] == selected_file and (selected_page == 0 or pdf["page"] == selected_page):
                        text_to_summarize += pdf["text"] or ""
            if text_to_summarize.strip():
                formatted = prompt.format(text=text_to_summarize)
                summary = model.invoke(formatted)
                st.success(summary)
            else:
                st.warning("No text found for the selected file/page.")


    with st.sidebar:
        st.title("🔑 Google API Key")
        api_key_input = st.text_input("Enter Google API Key", type="password", placeholder="GOOGLE_API_KEY")
        if st.button("Save API Key"):
            if api_key_input and api_key_input.strip():
                st.session_state["GOOGLE_API_KEY"] = api_key_input.strip()
                ensure_api_key()
                st.success("API key saved and configured.")
            else:
                st.warning("Please enter a valid API key.")
        st.write("---")
        st.title("📐 Embedding Provider")
        st.session_state["EMBEDDING_PROVIDER"] = st.selectbox(
            "Choose provider for embeddings",
            ["Google Gemini", "Local (No-Quota)"]
        )
        st.caption("Use Local if you hit quotas or 429 errors.")
        st.write("---")
        st.image("img/Robot.jpg")
        st.write("---")
        st.title("📁 File Upload Section")
        pdf_docs = st.file_uploader("Upload your PDF, TXT, or DOCX Files & \n Click on the Submit & Process Button ", accept_multiple_files=True, type=["pdf", "txt", "docx"])
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text, pdf_texts, tags, entities, summary, toc = get_pdf_text(pdf_docs)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.session_state.pdf_texts = pdf_texts
                st.session_state.raw_text = raw_text
                st.session_state.tags = tags
                st.session_state.entities = entities
                st.session_state.summary = summary
                st.session_state.toc = toc
                st.success("Done")
    # --- Comprehensive Document Knowledge Section ---
    st.markdown("---")
    st.subheader("📖 Document Knowledge & Analysis")
    if st.session_state.raw_text:
        st.markdown("**Summary:**")
        st.write(st.session_state.summary)
        st.markdown("**Table of Contents:**")
        for item in st.session_state.toc:
            st.write(f"- {item}")
        st.markdown("**Key Topics/Entities:**")
        st.write(", ".join(st.session_state.entities))
        # --- Knowledge Graph Visualization ---
        st.markdown("**Knowledge Graph:**")
        G = nx.Graph()
        for tag in st.session_state.tags:
            G.add_node(tag, color='gold')
        for entity in st.session_state.entities:
            G.add_node(entity, color='skyblue')
            for tag in st.session_state.tags:
                if tag in entity:
                    G.add_edge(tag, entity)
        pos = nx.spring_layout(G)
        fig, ax = plt.subplots(figsize=(8, 5))
        colors = [G.nodes[n].get('color', 'gray') for n in G.nodes]
        nx.draw(G, pos, with_labels=True, node_color=colors, ax=ax, font_size=10)
        st.pyplot(fig)
        st.write("---")
        st.image("img/pic.jpg")
        st.write("AI App created by @ Abhishek Kumar")
    # --- Analytics & Insights ---
    st.markdown("---")
    st.subheader("📊 Analytics & Insights & Auto-Tagging")
    # Show extracted tags and entities for navigation/filtering
    if st.session_state.tags or st.session_state.entities:
        st.markdown("**Extracted Tags:**")
        st.write(", ".join(st.session_state.tags) if st.session_state.tags else "No tags found.")
        st.markdown("**Extracted Entities:**")
        st.write(", ".join(st.session_state.entities) if st.session_state.entities else "No entities found.")
        # Tag-based filtering
        selected_tag = st.selectbox("Filter documents by tag", ["All"] + st.session_state.tags if st.session_state.tags else ["All"])
        if selected_tag != "All":
            filtered_docs = [pdf for pdf in st.session_state.pdf_texts if selected_tag in pdf["text"]]
            st.markdown(f"**Documents containing tag '{selected_tag}':**")
            for pdf in filtered_docs:
                st.write(f"{pdf['filename']} (Page {pdf['page']})")
    # Existing analytics
    if st.session_state.chat_history:
        # Most asked questions
        q_counts = {}
        for c in st.session_state.chat_history:
            q = c['question']
            q_counts[q] = q_counts.get(q, 0) + 1
        most_asked = sorted(q_counts.items(), key=lambda x: x[1], reverse=True)[:3]
        st.markdown("**Most Asked Questions:**")
        for q, count in most_asked:
            st.write(f"{q} ({count} times)")
        # Document coverage
        doc_counts = {}
        for c in st.session_state.chat_history:
            for src in c['source'].split(','):
                doc = src.split(' (Page ')[0].strip()
                if doc:
                    doc_counts[doc] = doc_counts.get(doc, 0) + 1
        st.markdown("**Document Coverage (Questions per Document):**")
        for doc, count in doc_counts.items():
            st.write(f"{doc}: {count} questions")
        # Confidence score (if available)
        st.markdown("**Answer Confidence:** N/A (not available from current model)")

    st.markdown(
        """
        <div style="position: fixed; bottom: 0; left: 0; width: 100%; background-color: #0E1117; padding: 15px; text-align: center;">
            © <a href="https://github.com/abhishekkumar62000" target="_blank">🧑‍💻Abhishek Kumar Yadav</a> | Made with ❤️
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()